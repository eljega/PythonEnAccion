"""
Traductor automatico de Documentos word con extension .docx

descripcion:
Este script traduce el texto de un documento word de un idioma a otro,
en cuestion de segundos.

Instrucciones de uso:
1. nos aseguramos de tener instaladas las dependencias: googletrans y python-docx con pip install googletrans==4.0.0-rc1 python-docx
2. debes modificar las variables ruta_archivo e idioma_destino segun el idioma que quieras y la ruta de tu archov word.
3.el documento traducido se guardara en la misma carpeta que el documento original.
"""

from googletrans import Translator, LANGUAGES
from docx import Document
import os

ruta_archivo = "ruta a la carpeta donde esta tu documento"  #aqui ingresas la ruta de tu archivo word
idioma_destino = "english"  #selecciona el idioma segun lo que requieras

def traducir_documento(ruta_archivo, idioma_destino):
    if idioma_destino not in LANGUAGES.values():
        print(f"Error: Idioma '{idioma_destino}' no es valido.")
        return
    try:
        documento = Document(ruta_archivo)
    except Exception as e:
        print(f"No se pudo abrir el archivo {ruta_archivo}: {e}")
        return

    traductor = Translator()
    for p in documento.paragraphs:
        if p.text:
            traduccion = traductor.translate(p.text, dest=idioma_destino).text
            p.clear()
            p.add_run(traduccion)
    directorio, nombre_archivo = os.path.split(ruta_archivo)
    nombre_archivo_traducido = f"traducido_{nombre_archivo}"
    ruta_archivo_traducido = os.path.join(directorio, nombre_archivo_traducido)
    documento.save(ruta_archivo_traducido)
    print(f"Documento traducido guardado como '{ruta_archivo_traducido}'")

traducir_documento(ruta_archivo, idioma_destino)
